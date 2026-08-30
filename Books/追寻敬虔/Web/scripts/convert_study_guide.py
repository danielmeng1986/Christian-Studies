#!/usr/bin/env python3
"""Convert the Word-normalized Puritan Studies guide DOCX to auditable Markdown."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
from typing import Any, Iterable
from zipfile import ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": WORD_NS, "r": REL_NS}
HEADING_STYLES = {f"Heading {level}": level for level in range(1, 5)}
BODY_STYLES = {"Normal", "Normal Indent", "gap"}


class ConversionError(ValueError):
    """Raised when the normalized Word file violates the conversion contract."""


def markdown_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("*", "\\*").replace("_", "\\_")


def merge_segments(segments: Iterable[tuple[bool, bool, str]]) -> list[tuple[bool, bool, str]]:
    source = list(segments)
    for index in range(1, len(source) - 1):
        bold, italic, text = source[index]
        if text.isspace() and source[index - 1][:2] == source[index + 1][:2]:
            source[index] = (*source[index - 1][:2], text)
    merged: list[tuple[bool, bool, str]] = []
    for bold, italic, text in source:
        if not text:
            continue
        if merged and merged[-1][:2] == (bold, italic):
            old_bold, old_italic, old_text = merged[-1]
            merged[-1] = (old_bold, old_italic, old_text + text)
        else:
            merged.append((bold, italic, text))
    return merged


def format_segments(segments: Iterable[tuple[bool, bool, str]]) -> str:
    rendered: list[str] = []
    for bold, italic, text in merge_segments(segments):
        escaped = markdown_escape(text)
        marker = "***" if bold and italic else "**" if bold else "*" if italic else ""
        rendered.append(f"{marker}{escaped}{marker}" if marker and text.strip() else escaped)
    return "".join(rendered)


def run_text(run_element: Any, footnote_numbers: dict[str, int]) -> str:
    values: list[str] = []
    for child in run_element:
        if child.tag == qn("w:t"):
            values.append(child.text or "")
        elif child.tag == qn("w:tab"):
            values.append("\t")
        elif child.tag in {qn("w:br"), qn("w:cr")}:
            values.append("\n")
        elif child.tag == qn("w:footnoteReference"):
            footnote_id = child.get(qn("w:id"))
            if footnote_id not in footnote_numbers:
                raise ConversionError(f"unknown footnote reference: {footnote_id}")
            number = footnote_numbers[footnote_id]
            values.append(f"[{number}](#脚注-{number})")
    return "".join(values)


def render_runs(paragraph: Paragraph, run_elements: Iterable[Any], footnote_numbers: dict[str, int]) -> str:
    segments: list[tuple[bool, bool, str]] = []
    for element in run_elements:
        run = Run(element, paragraph)
        segments.append((run.bold is True, run.italic is True, run_text(element, footnote_numbers)))
    return format_segments(segments)


def run_segment(paragraph: Paragraph, element: Any, footnote_numbers: dict[str, int]) -> tuple[bool, bool, str]:
    run = Run(element, paragraph)
    return run.bold is True, run.italic is True, run_text(element, footnote_numbers)


def render_paragraph(paragraph: Paragraph, footnote_numbers: dict[str, int]) -> str:
    parts: list[str] = []
    segments: list[tuple[bool, bool, str]] = []

    def flush() -> None:
        if segments:
            parts.append(format_segments(segments))
            segments.clear()

    for child in paragraph._p:
        if child.tag == qn("w:r"):
            segments.append(run_segment(paragraph, child, footnote_numbers))
        elif child.tag == qn("w:hyperlink"):
            flush()
            text = render_runs(paragraph, child.iter(qn("w:r")), footnote_numbers)
            relationship_id = child.get(qn("r:id"))
            if relationship_id and relationship_id in paragraph.part.rels:
                target = paragraph.part.rels[relationship_id].target_ref
                parts.append(f"[{text}]({target})")
            else:
                parts.append(text)
        else:
            segments.extend(
                run_segment(paragraph, element, footnote_numbers)
                for element in child.iter(qn("w:r"))
            )
    flush()
    # A leading manual break in Word is page layout, not paragraph content.
    return "".join(parts).replace("\r", "").lstrip("\t\n").rstrip()


def iter_blocks(parent: DocumentObject | _Cell) -> Iterable[Paragraph | Table]:
    parent_element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def table_markdown(table: Table, footnote_numbers: dict[str, int]) -> list[str]:
    rows: list[list[str]] = []
    seen_cells: set[int] = set()
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_identity = id(cell._tc)
            if cell_identity in seen_cells:
                cells.append("")
                continue
            seen_cells.add(cell_identity)
            content = [render_paragraph(paragraph, footnote_numbers) for paragraph in cell.paragraphs]
            cells.append("<br>".join(item for item in content if item).replace("|", "\\|"))
        rows.append(cells)
    if not rows:
        return []
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    return [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
        *("| " + " | ".join(row) + " |" for row in rows[1:]),
    ]


def load_footnotes(path: Path, document: DocumentObject) -> tuple[dict[str, int], dict[str, list[str]]]:
    with ZipFile(path) as archive:
        root = parse_xml(archive.read("word/footnotes.xml"))
    footnotes = [
        element
        for element in root.xpath("./w:footnote", namespaces=NS)
        if int(element.get(qn("w:id"))) > 0
    ]
    reference_ids = document.element.xpath(".//w:footnoteReference/@w:id")
    ordered_ids: list[str] = []
    for footnote_id in reference_ids:
        if footnote_id not in ordered_ids:
            ordered_ids.append(footnote_id)
    numbers = {footnote_id: index for index, footnote_id in enumerate(ordered_ids, 1)}
    bodies: dict[str, list[str]] = {}
    by_id = {element.get(qn("w:id")): element for element in footnotes}
    for footnote_id in ordered_ids:
        element = by_id.get(footnote_id)
        if element is None:
            raise ConversionError(f"footnote body missing: {footnote_id}")
        paragraphs: list[str] = []
        for paragraph_element in element.xpath("./w:p", namespaces=NS):
            paragraph = Paragraph(paragraph_element, document.part)
            text = render_paragraph(paragraph, numbers).strip()
            if text:
                paragraphs.append(text)
        bodies[footnote_id] = paragraphs
    if len(numbers) != len(footnotes):
        raise ConversionError("unreferenced or duplicated Word footnotes detected")
    return numbers, bodies


def convert(docx_path: Path, source_path: Path, pages: str) -> tuple[str, dict[str, Any]]:
    document = Document(docx_path)
    footnote_numbers, footnotes = load_footnotes(docx_path, document)
    source_hash = hashlib.sha256(source_path.read_bytes()).hexdigest()
    nonempty_styles = {
        paragraph.style.name
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    }
    unsupported = nonempty_styles - set(HEADING_STYLES) - BODY_STYLES
    if unsupported:
        raise ConversionError(f"unsupported paragraph styles: {sorted(unsupported)}")

    lines = [
        "---",
        "title: 清教徒研究的教牧應用",
        "subtitle: Pastoral Application of Puritan Studies",
        "author: Paul Chang（張麟至）",
        "document_type: study_guide",
        f"source_file: {json.dumps(source_path.name, ensure_ascii=False)}",
        f"source_location: {json.dumps(f'Microsoft Word 文档物理第{pages}页', ensure_ascii=False)}",
        f"source_sha256: {source_hash}",
        "language: zh-Hant",
        "---",
        "",
        f"> 来源定位：原始文档 `{source_path.name}`，Microsoft Word 文档物理第{pages}页。",
        "",
    ]
    style_counts: dict[str, int] = {}
    paragraph_count = 0
    table_count = 0
    quote_open = False
    for block in iter_blocks(document):
        if isinstance(block, Table):
            quote_open = False
            table_count += 1
            lines.extend(table_markdown(block, footnote_numbers))
            lines.append("")
            continue
        text = render_paragraph(block, footnote_numbers)
        if not text:
            continue
        paragraph_count += 1
        style = block.style.name
        style_counts[style] = style_counts.get(style, 0) + 1
        if style in HEADING_STYLES:
            quote_open = False
            lines.extend([f"{'#' * HEADING_STYLES[style]} {text}", ""])
        elif style == "Normal Indent":
            if quote_open:
                lines.append(">")
            lines.extend([f"> {text}", ""])
            quote_open = True
        else:
            quote_open = False
            lines.extend([text, ""])

    lines.extend(["# 脚注", ""])
    for footnote_id, number in footnote_numbers.items():
        lines.extend([f"### 脚注 {number}", ""])
        for paragraph in footnotes[footnote_id]:
            lines.extend([paragraph, ""])

    markdown = "\n".join(lines).rstrip() + "\n"
    audit = {
        "sourceSha256": source_hash,
        "sourcePages": pages,
        "sourceParagraphs": len(document.paragraphs),
        "convertedNonemptyBodyParagraphs": paragraph_count,
        "tables": table_count,
        "footnoteReferences": len(document.element.xpath(".//w:footnoteReference")),
        "footnoteBodies": len(footnotes),
        "styleCounts": style_counts,
        "mappedNonSemanticStyles": {"gap": style_counts.get("gap", 0)},
    }
    return markdown, audit


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path, help="DOCX copy saved by Microsoft Word")
    parser.add_argument("--source", type=Path, required=True, help="Original legacy Word source")
    parser.add_argument("--source-pages", required=True, help="Verified Microsoft Word page range")
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--audit", type=Path)
    args = parser.parse_args()
    markdown, audit = convert(args.docx, args.source, args.source_pages)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(markdown, encoding="utf-8")
    if args.audit:
        args.audit.parent.mkdir(parents=True, exist_ok=True)
        args.audit.write_text(json.dumps(audit, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
